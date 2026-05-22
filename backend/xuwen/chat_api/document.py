"""文档文本提取。

按扩展名 dispatch 到不同 extractor。所有 extractor 把字节流转为 markdown 风格纯文本。
失败时抛 DocumentError；过大或不支持的格式也走它。

支持格式：
- 纯文本：txt / md / json / csv / log / yml / yaml / xml / ini
- PDF：pypdf
- DOCX：python-docx
- XLSX：openpyxl（按 sheet 输出 markdown 表格）
- HTML：beautifulsoup4
"""

from __future__ import annotations

import csv
import io
import json
import logging
from collections.abc import Callable
from dataclasses import dataclass

from xuwen.core.errors import XuwenError

logger = logging.getLogger(__name__)


class DocumentError(XuwenError):
    """文档提取失败。"""

    code = "xuwen.document"
    http_status = 400


@dataclass(slots=True, frozen=True)
class ExtractedDocument:
    """提取结果。"""

    filename: str
    extension: str
    text: str
    char_count: int
    # 粗略 token 估算（中文 ~1.5 char/token, 英文 ~4 char/token）
    estimated_tokens: int


# ---------------------------------------------------------------------------
# 公开接口
# ---------------------------------------------------------------------------


_DEFAULT_MAX_BYTES = 16 * 1024 * 1024  # 16MB


def extract(
    data: bytes,
    filename: str,
    *,
    max_bytes: int = _DEFAULT_MAX_BYTES,
) -> ExtractedDocument:
    """从字节流提取文本。"""
    if len(data) > max_bytes:
        mb = max_bytes / (1024 * 1024)
        raise DocumentError(f"文件过大（>{mb:.0f}MB），请压缩或切分后再上传")

    ext = _ext_of(filename)
    extractor = _DISPATCH.get(ext)
    if extractor is None:
        raise DocumentError(
            f"暂不支持 .{ext or '?'} 格式。已支持：{sorted(_DISPATCH.keys())}"
        )

    try:
        text = extractor(data)
    except DocumentError:
        raise
    except Exception as e:
        # 不带原始错误信息（防泄漏）
        logger.warning("文档提取失败 ext=%s: %s", ext, type(e).__name__)
        raise DocumentError(
            f"无法解析 {filename}（{type(e).__name__}）。可能是文件损坏或格式不被支持。"
        ) from e

    text = text.strip()
    if not text:
        raise DocumentError(f"{filename} 提取后内容为空")

    char_count = len(text)
    return ExtractedDocument(
        filename=filename,
        extension=ext,
        text=text,
        char_count=char_count,
        estimated_tokens=_estimate_tokens(text),
    )


# ---------------------------------------------------------------------------
# extractors
# ---------------------------------------------------------------------------


def _extract_plain(data: bytes) -> str:
    return _decode(data)


def _extract_csv(data: bytes) -> str:
    text = _decode(data)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    lines: list[str] = []
    # 把 CSV 转成 markdown 表格（前 200 行，避免太长）
    head = rows[0]
    lines.append("| " + " | ".join(head) + " |")
    lines.append("| " + " | ".join("---" for _ in head) + " |")
    for row in rows[1 : 200 + 1]:
        # 列数不齐时补齐
        padded = (row + [""] * len(head))[: len(head)]
        lines.append("| " + " | ".join(padded) + " |")
    if len(rows) > 201:
        lines.append(f"\n（共 {len(rows)} 行，仅展示前 200 行）")
    return "\n".join(lines)


def _extract_json(data: bytes) -> str:
    try:
        obj = json.loads(_decode(data))
    except json.JSONDecodeError as e:
        raise DocumentError(f"JSON 解析失败：{e.msg} (行 {e.lineno})") from e
    return json.dumps(obj, ensure_ascii=False, indent=2)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, 1):
        try:
            txt = page.extract_text() or ""
        except Exception as e:
            logger.info("pdf 第 %d 页提取失败：%s", i, type(e).__name__)
            continue
        if txt.strip():
            parts.append(f"--- 第 {i} 页 ---\n{txt.strip()}")
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows_text.append("| " + " | ".join(cells) + " |")
        if rows_text:
            parts.append("\n".join(rows_text))
    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet：{sheet_name}\n")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        head = [_xcell(c) for c in rows[0]]
        parts.append("| " + " | ".join(head) + " |")
        parts.append("| " + " | ".join("---" for _ in head) + " |")
        for row in rows[1 : 200 + 1]:
            cells = [_xcell(c) for c in row]
            padded = (cells + [""] * len(head))[: len(head)]
            parts.append("| " + " | ".join(padded) + " |")
        if len(rows) > 201:
            parts.append(f"（共 {len(rows)} 行，仅展示前 200 行）")
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(_decode(data), "html.parser")
    # 去掉脚本和样式
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return text


# 扩展名 → 提取器
_DISPATCH: dict[str, Callable[[bytes], str]] = {
    "txt": _extract_plain,
    "md": _extract_plain,
    "log": _extract_plain,
    "yml": _extract_plain,
    "yaml": _extract_plain,
    "xml": _extract_plain,
    "ini": _extract_plain,
    "csv": _extract_csv,
    "json": _extract_json,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "html": _extract_html,
    "htm": _extract_html,
}


def supported_extensions() -> list[str]:
    return sorted(_DISPATCH.keys())


# ---------------------------------------------------------------------------
# 内部
# ---------------------------------------------------------------------------


def _ext_of(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def _decode(data: bytes) -> str:
    """尝试常见编码解码（覆盖中文场景）。

    顺序：带 BOM 的 utf-16 → utf-8 → gbk → gb18030 → big5 → 兜底。
    把无 BOM 的 utf-16 推后，因为 GBK 文本经常会被它"成功"解码成乱码。
    """
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        try:
            return data.decode("utf-16")
        except UnicodeDecodeError:
            pass
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "big5"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _xcell(value: object) -> str:
    if value is None:
        return ""
    return str(value).replace("|", "\\|").replace("\n", " ").strip()


def _estimate_tokens(text: str) -> int:
    """粗略估算 token 数。中文按 1.5 char/token、英文按 4 char/token 混合估。"""
    if not text:
        return 0
    cn_chars = sum(1 for c in text if "一" <= c <= "鿿")
    en_chars = len(text) - cn_chars
    return int(cn_chars / 1.5 + en_chars / 4.0)
